import json, subprocess, os, sys
SK = "/Users/huynguyen/Library/Application Support/Claude/local-agent-mode-sessions/skills-plugin/f04a2c68-340a-4e46-9ccd-208841bc29ff/bd4055f6-5961-407c-88ef-e871fa98c8c8/skills/humiley-brand"
OUT = "/Users/huynguyen/Library/CloudStorage/OneDrive-Humiley(2)/Claude Projects/TimeKeeping Web App/deliverables/Humiley-Portal-User-Manual.docx"
SHOTS = "/tmp/shots/"

def fig(name): return {"type": "paragraph", "text": "[[FIG:%s]]" % name}
def lab(t): return {"type": "label", "text": t}
def para(t): return {"type": "paragraph", "text": t}
def bl(items): return {"type": "bullets", "items": items}
def num(items): return {"type": "numbers", "items": items}

sections = [
 {"no": "02", "title": "Signing in & access levels", "blocks": [
   lab("HOW TO SIGN IN"),
   num(["Open the portal URL in your browser.",
        "Click “Sign in with Microsoft 365” and choose your Humiley account.",
        "You land on the dashboard for your access level."]),
   para("Access is governed by five cumulative levels — each can do everything the level before it can, plus more. Your administrator sets your level under Access & Permissions."),
   {"type": "table", "header": ["Level", "Can do"], "rows": [
     ["User", "Self-service: own profile, attendance, leave, claims, training, plus CRM & assigned projects."],
     ["Contributor", "+ approve team requests, People/HR modules, and create projects."],
     ["Approver", "+ view Payroll & Finance."],
     ["Editor", "+ run / edit Payroll."],
     ["Admin", "+ Access & Permissions (assign levels)."]]},
   fig("04_permissions")]},
 {"no": "03", "title": "Your dashboard & self-service", "blocks": [
   para("The dashboard is a live snapshot of the organisation — headcount, attendance, leave, recruitment and performance — with PowerBI-style charts that update from real data and a department filter. Export the current view to a branded PDF at any time."),
   bl(["Check in / out and view your attendance.", "Request leave, submit travel-expense claims, set goals and complete training.",
       "Everything you submit routes to your direct manager for approval."]),
   fig("01_company_dashboard")]},
 {"no": "04", "title": "People & HR", "blocks": [
   para("Managers and HR maintain the full employee record — role, department, assigned devices and documents — and run recruitment, onboarding, performance (PADR) and talent reviews."),
   bl(["Add or edit employees; import a batch from Excel/CSV; export for reporting.",
       "Onboarding checklists sync to the employee's My Training.",
       "Sensitive HR data (reviews, talent, exits) is visible to managers only."]),
   fig("02_employee_database")]},
 {"no": "05", "title": "Payroll & finance", "blocks": [
   para("Payroll computes monthly pay runs with personal income tax, social insurance and employer cost automatically, and keeps history from 2024 for trend analysis."),
   bl(["Viewing payroll requires Approver level or above.",
       "Only Editor / Admin can run or edit a pay run.",
       "Every figure is exportable to a branded report."]),
   fig("03_payroll")]},
 {"no": "06", "title": "PMC — project portfolio & workspace", "blocks": [
   para("The PMC suite manages projects to PMBOK practice. The portfolio shows every project at a glance — contract value, health (RAG), CPI/SPI and milestones. Open a project to enter its workspace, organised into tabs."),
   para("Staff see only the projects they are assigned to; managers see the whole portfolio."),
   fig("05_pmc_portfolio")]},
 {"no": "07", "title": "Project Overview", "blocks": [
   para("The Overview opens with KPI tiles and PowerBI-style charts — cost (EVM), deliverables by status, schedule tasks and risk exposure — followed by the project charter, change-control log and lessons learned."),
   fig("06_pm_overview")]},
 {"no": "08", "title": "Schedule (Gantt)", "blocks": [
   para("The Schedule is a Gantt timeline showing tasks, milestones, the critical path and a “today” marker. Filter by all time, a year, a month, or a custom day-to-day range, and export the schedule to PDF."),
   fig("07_pm_schedule")]},
 {"no": "09", "title": "Cost & Earned Value (EVM)", "blocks": [
   para("Cost tracks budget vs committed vs actual by category and computes earned-value metrics — CPI, SPI and the forecast at completion (EAC) — to give an early warning on cost and schedule performance."),
   fig("08_pm_cost_evm")]},
 {"no": "10", "title": "Risk register & heatmap", "blocks": [
   para("Risks are scored on a probability × impact heatmap and banded automatically from Low to Critical. The register can be filtered by category, response strategy, owner, contractor or status."),
   fig("09_pm_risk")]},
 {"no": "11", "title": "Quality & Inspection-Test Plans", "blocks": [
   para("The quality register captures inspections, NCRs, audits and tests. Inspection & Test Plans (ITP) appear on a timeline with planned start/finish dates and a day-range filter."),
   fig("10_pm_quality_itp")]},
 {"no": "12", "title": "Procurement, contracts & stakeholders", "blocks": [
   para("Procurement records subcontract packages, vendors, values, retention and interim payment certificates (IPC). Cost and procurement edits are restricted to managers."),
   fig("11_pm_procurement"),
   para("Stakeholders are mapped on a Mendelow power / interest grid so the team can plan engagement by quadrant, tracking influence, interest and attitude."),
   fig("12_pm_stakeholders")]},
 {"no": "13", "title": "Documents (SharePoint)", "blocks": [
   lab("BUILD THE PROJECT FOLDERS"),
   num(["Open the project, click Edit, and paste the project's SharePoint document-library link into “SharePoint docs folder”.",
        "Click Save & build folders — when signed in with Microsoft 365, the standard 50-folder PMC structure is created in SharePoint automatically.",
        "Browse every folder in the Documents tab; each is a one-click link straight into SharePoint."]),
   para("The folder list carries owner, confidentiality and retention for each folder. SharePoint permissions are managed by IT."),
   fig("13_pm_documents_sharepoint")]},
 {"no": "14", "title": "Sales pipeline (CRM)", "blocks": [
   para("CRM tracks deals by stage with value and probability, alongside companies, contacts and products. A Won deal can be converted into a PMC project in a single click (Contributor level or above)."),
   fig("14_crm_pipeline")]},
 {"no": "15", "title": "Working in English or Vietnamese", "blocks": [
   para("Switch the entire portal between English and Vietnamese using the EN / VN flag in the top bar. Labels, menus, tables and chart legends all translate."),
   fig("15_vietnamese_dashboard")]},
 {"no": "16", "title": "English – Vietnamese glossary", "blocks": [
   {"type": "table", "header": ["English", "Tiếng Việt"], "rows": [
     ["Dashboard", "Bảng điều khiển"], ["Attendance", "Chấm công"], ["Leave", "Nghỉ phép"],
     ["Payroll", "Lương"], ["Project", "Dự án"], ["Schedule", "Tiến độ"], ["Cost / EVM", "Chi phí / EVM"],
     ["Risk", "Rủi ro"], ["Quality", "Chất lượng"], ["Procurement", "Mua sắm"], ["Stakeholders", "Bên liên quan"],
     ["Documents", "Tài liệu"], ["Pipeline", "Quy trình bán hàng"], ["Access & Permissions", "Phân quyền"]]}]},
 {"no": "17", "title": "Administration & data safety", "blocks": [
   bl(["Updates replace only the application code — the database (timekeeping.db) is never overwritten, so all data is preserved.",
       "Keep the database on a persistent volume and back it up before each update.",
       "Access levels and module access are managed under Access & Permissions; changes apply at the user's next sign-in."])]},
]

spec = {
  "doc_type": "User Manual", "title": "People & Workplace", "title2": "Portal.",
  "oneliner": "End-user guide — HR, CRM and PMC project management",
  "running_title": "Humiley Portal — User Manual", "doc_code": "HML-MAN-0001  Rev 01.0",
  "document_number": "HML-MAN-0001", "issue_date": "2026", "owner": "Humiley Group Inc.",
  "approval": "Tony Nguyen, CEO", "distribution": "All staff", "confidentiality": "Internal",
  "about": "This manual is a practical, screen-by-screen guide to the Humiley People & Workplace Portal. Screenshots are taken from the live demonstration dataset. It covers signing in, access levels, HR self-service, the PMC project workspace, CRM, payroll and administration.",
  "sections": sections,
}
json.dump(spec, open("/tmp/build/manual_spec.json", "w"))
r = subprocess.run([sys.executable, SK + "/scripts/fill_document.py", "--template", "EN",
                    "--out", OUT, "--spec", "/tmp/build/manual_spec.json"], capture_output=True, text=True)
print("fill:", r.returncode, (r.stderr or r.stdout)[-400:])

# ---- post-process: replace [[FIG:name]] markers with the screenshot ----
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
doc = Document(OUT); n = 0
for p in doc.paragraphs:
    t = p.text.strip()
    if t.startswith("[[FIG:") and t.endswith("]]"):
        name = t[6:-2]
        for r2 in list(p.runs): r2.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        path = SHOTS + name + ".jpg"
        if os.path.exists(path):
            p.add_run().add_picture(path, width=Inches(6.3)); n += 1
        else:
            p.add_run("[missing figure: %s]" % name)
doc.save(OUT)
print("inserted %d screenshots; saved %s" % (n, OUT))
